from PIL import Image
from docx import Document
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import email
from email import policy
from email.message import EmailMessage
from email.parser import BytesParser

import requests
from langchain_community.document_loaders import *

from werkzeug.utils import secure_filename
import zipfile
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
import shutil

from pdf2image import convert_from_path
import pytesseract
from langchain.schema.document import Document

supported_formats = (
    '.txt', '.pdf', '.docx', '.png', '.jpg',
    ".doc", ".ppt", ".pptx", ".xls", ".xlsx",
    ".htm", ".html", ".eml", ".csv"
)


# 解压缩包
def unzip(zip_file_path, output_folder):
    # 确保目标文件夹存在，不存在则创建
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 打开ZIP文件
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        # 遍历ZIP文件中的所有内容
        for member in zip_ref.infolist():
            # 如果是文件而不是文件夹
            if not member.is_dir():
                # 获取文件的最终文件名（不包含子文件夹路径）
                extracted_file_name = os.path.basename(member.filename)
                if extracted_file_name:  # 确保不是空字符串
                    # 创建完整的保存路径
                    extracted_file_path = os.path.join(output_folder, extracted_file_name)

                    # 处理文件名冲突：如果文件已存在，重新命名或覆盖
                    counter = 1
                    while os.path.exists(extracted_file_path):
                        name, ext = os.path.splitext(extracted_file_name)
                        extracted_file_path = os.path.join(output_folder, f"{name}_{counter}{ext}")
                        counter += 1

                    # 将文件提取到目标路径
                    with zip_ref.open(member) as source_file:
                        with open(extracted_file_path, 'wb') as target_file:
                            shutil.copyfileobj(source_file, target_file)

    return output_folder


# 压缩包
def save_translated_zip(translated_unzipped_files):
    # 指定压缩包的路径和名称
    output_zip_path = "static/translated/output.zip"

    # 打开一个新创建的zip文件进行写入操作
    with zipfile.ZipFile(output_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # 遍历所有文件并添加到压缩包中
        for file in translated_unzipped_files:
            zipf.write(file)

    print(f"压缩完成，文件已保存至 {output_zip_path}")
    return output_zip_path


# 获取文件类型
def get_file_type(file_path: str):
    # 获取文件类型
    _, file_type = os.path.splitext(file_path)
    file_type = file_type.lower()
    return file_type


# 判断文件类型
def is_supported_formats(file_path: str):
    file_type = get_file_type(file_path)
    if file_type not in supported_formats:
        return False
    else:
        return True


# 导入eml文件
def import_email(file_path: str, base_output_folder: str):
    # 读取 .eml/.msg 文件
    with open(file_path, 'rb') as f:
        msg = BytesParser(policy=policy.default).parse(f)

    if not os.path.exists(base_output_folder):
        os.makedirs(base_output_folder)

    # 创建文件夹
    email_id = os.path.splitext(os.path.basename(file_path))[0]
    output_folder = os.path.join(base_output_folder, email_id)

    # 获取邮件头信息并保存到字典中
    email_headers = {
        "From": msg['From'],
        "To": msg['To'],
        "Cc": msg.get('Cc', 'None'),
        "Subject": msg.get('Subject', 'None'),
        # 获取邮件发送时间
        "Date": msg.get('Date', 'None')
    }

    # 生成邮件头信息的字符串格式
    header_info = (f"发件人：{email_headers['From']}\n"
                   f"收件人：{email_headers['To']}\n"
                   f"抄送：{email_headers['Cc']}\n"
                   f"日期：{email_headers['Date']}\n\n"
                   f"正文：")

    # 获取邮件正文信息
    doc = UnstructuredEmailLoader(file_path=file_path, process_attachments=False).load()

    # 获取邮件附件，并创建邮件文件夹
    save_email_attachments(file_path, output_folder)

    # 将邮件头信息添加到正文内容开头
    doc[0].page_content = header_info + doc[0].page_content

    # 返回处理后的 doc 和 email_headers
    return {"doc": doc, "email_headers": email_headers, "output_folder": output_folder}


# 导入文件
def import_file(file_path: str, lang='osd'):
    # 获取文件类型
    _, file_type = os.path.splitext(file_path)
    file_type = file_type.lower()

    if file_type not in supported_formats:
        return False

    elif file_type in '.pdf':
        # doc = PyPDFLoader(file_path, extract_images=True).load()
        # num_pages = len(doc)
        # for i in range(1, num_pages):
        #     doc[0].page_content = doc[0].page_content + doc[i].page_content
        """将 PDF 文件 OCR 处理后保存为文本文件"""
        # Step 1: 将 PDF 转换为图像
        images = convert_from_path(file_path, dpi=300)
        # Step 2: 初始化一个空字符串来存储 OCR 结果
        full_text = ""

        # Step 3: 对每个图像进行 OCR 识别
        for i, image in enumerate(images):
            print(f"Processing page {i + 1}/{len(images)}...")
            text = pytesseract.image_to_string(image, lang)
            full_text += text + "\n"

        doc = [Document(page_content=full_text)]

    elif file_type in ('.png', '.jpg', 'jpeg'):
        # doc = UnstructuredImageLoader(file_path).load()
        # 打开图像文件
        img = Image.open(file_path)

        # 使用 pytesseract 进行 OCR 识别
        text = pytesseract.image_to_string(img, lang)
        doc = [Document(page_content=text)]

    elif file_type in (".doc", ".docx"):
        doc = Docx2txtLoader(file_path).load()

    elif file_type in (".ppt", ".pptx"):
        doc = UnstructuredPowerPointLoader(file_path).load()

    elif file_type in (".xls", ".xlsx"):
        doc = UnstructuredExcelLoader(file_path).load()

    elif file_type in (".htm", ".html"):
        doc = BSHTMLLoader(file_path, open_encoding="unicode_escape").load()

    elif file_type in ('.eml', '.msg'):
        doc = UnstructuredEmailLoader(file_path, process_attachments=True).load()

    elif file_type in ".csv":
        doc = CSVLoader(file_path).load()

    elif file_type in ".txt":
        doc = TextLoader(file_path, autodetect_encoding=True).load()

    else:
        doc = TextLoader(file_path, autodetect_encoding=True).load()

    return doc


# 保存eml及其附件于一个文件夹
def save_email_attachments(eml_file_path, output_folder):
    # 创建输出文件夹
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 使用不同的变量名来打开文件，避免变量冲突
    with open(eml_file_path, 'rb') as eml_file:
        msg = email.message_from_binary_file(eml_file)

    # 遍历邮件的各个部分
    for part in msg.walk():
        # 如果邮件内容是附件
        if part.get_content_maintype() == "multipart":
            continue
        if part.get("Content-Disposition") is None:
            continue

        # 获取附件文件名
        filename = part.get_filename()
        if filename:
            # 解码附件文件名并确保安全
            filename = email.header.decode_header(filename)[0][0]
            if isinstance(filename, bytes):
                filename = filename.decode()
            filename = secure_filename(filename)  # 确保文件名安全

            # 保存附件到指定文件夹
            attachment_path = os.path.join(output_folder, filename)
            with open(attachment_path, "wb") as f:
                f.write(part.get_payload(decode=True))

    return output_folder


# 获取文件夹中所有文件的绝对路径（不包括子文件夹）
def get_files_with_absolute_paths(directory):
    file_path = [os.path.abspath(os.path.join(directory, f)) for f in os.listdir(directory) if
                 os.path.isfile(os.path.join(directory, f))]
    return file_path


# 删除文件夹
def delete_folder(folder_path):
    # 检查文件夹是否存在
    if os.path.exists(folder_path):
        # 使用shutil.rmtree删除文件夹及其内容
        shutil.rmtree(folder_path)
    else:
        print(f"文件夹不存在: {folder_path}")


# 保存已翻译文件为EML
def save_translated_email(filepath, email_headers, translated_content, translated_folder_path,
                          translated_attachments=None):
    new_msg = EmailMessage()
    new_msg['From'] = email_headers["From"]
    new_msg['To'] = email_headers["To"]
    new_msg['Subject'] = email_headers['Subject']
    new_msg.set_content(translated_content)

    filename = os.path.basename(filepath)
    translated_filename = os.path.splitext(filename)[0] + '_Translated.eml'
    translated_filepath = os.path.join(translated_folder_path, translated_filename)

    if translated_attachments is not None:
        # 将翻译后的附件添加到新邮件中
        for docx_path in translated_attachments:
            with open(docx_path, 'rb') as f:
                new_msg.add_attachment(f.read(), maintype='application',
                                       subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                                       filename=os.path.basename(docx_path))

    # 保存翻译后的文件
    with open(translated_filepath, 'wb') as f:
        f.write(new_msg.as_bytes())

    return translated_filepath


# 保存已翻译文件为TXT
def save_translated_txt(filepath, translated_content, translated_folder_path):
    filename = os.path.basename(filepath)
    translated_filename = os.path.splitext(filename)[0] + '_Translated.txt'
    translated_filepath = os.path.join(translated_folder_path, translated_filename)

    with open(translated_filepath, 'w', encoding='utf-8') as txt_file:
        txt_file.write(translated_content)

    # 返回翻译后文件的路径
    return translated_filepath


# 保存已翻译文件为DOCX
def save_translated_docx(filepath, translated_content, translated_folder_path):
    filename = os.path.basename(filepath)
    translated_filename = os.path.splitext(filename)[0] + '_Translated.docx'
    translated_filepath = os.path.join(translated_folder_path, translated_filename)

    # 创建一个新的docx文件
    doc = Document()
    doc.add_paragraph(translated_content)

    # 保存为.docx文件
    doc.save(translated_filepath)

    # 返回翻译后文件的路径
    return translated_filepath


# 保存已翻译文件为PDF
def save_translated_pdf(filepath, translated_content, translated_folder_path):
    # 设置字体路径，确保下载并提供支持中文的字体
    font_path = "localbin/translator/SimSun.ttf"  # 替换为你下载的字体文件路径
    if not os.path.exists(font_path):
        raise FileNotFoundError(f"字体文件未找到：{font_path}")

    # 注册支持中文的字体
    pdfmetrics.registerFont(TTFont('SimSun', font_path))

    filename = os.path.basename(filepath)
    translated_filename = os.path.splitext(filename)[0] + '_Translated.pdf'
    translated_filepath = os.path.join(translated_folder_path, translated_filename)

    # 创建 PDF 文件
    pdf = canvas.Canvas(translated_filepath, pagesize=letter)
    width, height = letter

    # 设置字体为 SimSun，字号为 12
    pdf.setFont("SimSun", 12)

    # 将翻译内容分行添加到 PDF
    y_position = height - 50  # 初始Y位置
    line_height = 16  # 行高

    for line in translated_content.splitlines():
        pdf.drawString(50, y_position, line)
        y_position -= line_height
        if y_position < 50:  # 如果到达页面底部，创建新页面
            pdf.showPage()
            pdf.setFont("SimSun", 12)  # 新页面需要重新设置字体
            y_position = height - 50

    # 保存 PDF 文件
    pdf.save()

    # 返回翻译后文件的路径
    return translated_filepath


# 生成文件URL链接
def generate_file_url(file_path, base_url):
    # 获取文件名
    file_name = os.path.basename(file_path)
    translated_url = f"{base_url}/{file_name}"

    return translated_url
